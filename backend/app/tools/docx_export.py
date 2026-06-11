"""Server-side Word (DOCX) generation for research reports using python-docx."""
import io
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


_BLUE = RGBColor(0x1E, 0x40, 0xAF)
_GREY = RGBColor(0x6B, 0x72, 0x80)
_DARK = RGBColor(0x11, 0x18, 0x27)


def _set_heading_color(paragraph, color: RGBColor):
    for run in paragraph.runs:
        run.font.color.rgb = color


def _add_horizontal_rule(doc: Document):
    """Insert a thin blue horizontal line via XML."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E40AF")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_margins(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def generate_report_docx(report_dict: dict, session_id: str) -> bytes:
    """Generate a DOCX from a research report dict and return raw bytes."""
    doc = Document()
    _set_margins(doc)

    # ── Default body style ────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = _DARK

    topic = report_dict.get("topic", "Research Report")
    summary = report_dict.get("summary", "")
    sections = report_dict.get("sections", [])
    sources = report_dict.get("sources", [])
    kg = report_dict.get("knowledge_graph", {})
    entities = kg.get("entities", []) if isinstance(kg, dict) else []
    created_at = report_dict.get("created_at", "")

    # ── Title ─────────────────────────────────────────────────────────────────
    title_para = doc.add_heading(topic, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title_para.runs:
        run.font.color.rgb = _DARK
        run.font.size = Pt(24)
        run.font.bold = True
    _add_horizontal_rule(doc)

    # Metadata
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    meta.add_run("Research ID: ").bold = True
    meta.add_run(session_id)
    meta.add_run("   |   ").font.color.rgb = _GREY
    if created_at:
        meta.add_run("Generated: ").bold = True
        meta.add_run(created_at[:19].replace("T", " ") + " UTC")
    for run in meta.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = _GREY

    # ── Executive Summary ─────────────────────────────────────────────────────
    h = doc.add_heading("Executive Summary", level=1)
    _set_heading_color(h, _BLUE)
    _add_horizontal_rule(doc)

    if summary:
        for para in summary.split("\n\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(6)

    # ── Sections ──────────────────────────────────────────────────────────────
    for section in sections:
        heading = section.get("heading", "")
        content = section.get("content", "")
        if not heading and not content:
            continue
        h = doc.add_heading(heading, level=2)
        _set_heading_color(h, _BLUE)
        _add_horizontal_rule(doc)
        for para in content.split("\n\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(6)

    # ── Knowledge Graph Entities ──────────────────────────────────────────────
    if entities:
        doc.add_page_break()
        h = doc.add_heading("Key Concepts & Entities", level=1)
        _set_heading_color(h, _BLUE)
        _add_horizontal_rule(doc)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for cell, label in zip(hdr, ["Entity", "Type", "Description"]):
            cell.text = label
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Blue cell background
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1E40AF")
            tcPr.append(shd)

        for ent in entities[:30]:
            row = table.add_row().cells
            row[0].text = ent.get("name", "")
            row[1].text = ent.get("type", "")
            row[2].text = ent.get("description", "")[:150]
            for cell in row:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph()

    # ── Sources ───────────────────────────────────────────────────────────────
    if sources:
        doc.add_page_break()
        h = doc.add_heading("Sources & References", level=1)
        _set_heading_color(h, _BLUE)
        _add_horizontal_rule(doc)

        for i, src in enumerate(sources, 1):
            title = src.get("title", "Untitled") if isinstance(src, dict) else getattr(src, "title", "Untitled")
            url = src.get("url", "") if isinstance(src, dict) else getattr(src, "url", "")
            summary_text = src.get("summary", "") if isinstance(src, dict) else getattr(src, "summary", "")

            p = doc.add_paragraph()
            run = p.add_run(f"[{i}] {title}")
            run.bold = True
            run.font.color.rgb = _BLUE
            run.font.size = Pt(10)

            url_p = doc.add_paragraph(url)
            url_p.runs[0].font.color.rgb = _GREY
            url_p.runs[0].font.size = Pt(9)
            url_p.paragraph_format.space_before = Pt(0)

            if summary_text:
                s_p = doc.add_paragraph(summary_text[:250])
                s_p.runs[0].font.size = Pt(10)
                s_p.paragraph_format.space_after = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
